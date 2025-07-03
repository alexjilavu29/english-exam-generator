from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_word_document(exam_questions, title="English Exam", include_answers=True, format_options=None):
    """
    Create a formatted Word document with exam questions.
    
    Args:
        exam_questions: List of question dictionaries
        title: Title of the exam
        include_answers: Whether to include answer key
        format_options: Dictionary of formatting options
    
    Returns:
        Path to the created document
    """
    # Set default format options if not provided
    if format_options is None:
        format_options = {
            'font_name': 'Calibri',
            'font_size': 11,
            'title_font_size': 16,
            'heading_font_size': 14,
            'page_orientation': 'portrait',
            'question_spacing': 0,
            'include_header_footer': True,
            'include_question_type': True
        }
    
    doc = Document()
    
    # Set page orientation
    if format_options.get('page_orientation') == 'landscape':
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add header/footer if requested
    if format_options.get('include_header_footer'):
        add_header_footer(doc, title)
    
    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(format_options.get('title_font_size', 16))
    title_run.font.name = format_options.get('font_name', 'Calibri')
    
    # Add instructions
    instructions = doc.add_paragraph()
    instructions.alignment = WD_ALIGN_PARAGRAPH.LEFT
    instructions_run = instructions.add_run('Choose the correct answer for each question.')
    instructions_run.italic = True
    instructions_run.font.size = Pt(format_options.get('font_size', 11))
    instructions_run.font.name = format_options.get('font_name', 'Calibri')
    
    doc.add_paragraph()  # Add space
    
    # Group questions by topic if requested
    if format_options.get('include_question_type'):
        # Sort questions by topic
        vocab_questions = [q for q in exam_questions if q['topic'] == 'Vocabulary']
        grammar_questions = [q for q in exam_questions if q['topic'] == 'Grammar']
        
        # Add vocabulary section if there are vocabulary questions
        if vocab_questions:
            vocab_heading = doc.add_heading('Vocabulary Questions', level=1)
            vocab_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for heading_part in vocab_heading.runs:
                heading_part.font.size = Pt(format_options.get('heading_font_size', 14))
                heading_part.font.name = format_options.get('font_name', 'Calibri')
            
            add_questions_to_document(doc, vocab_questions, format_options, start_index=1)
        
        # Add grammar section if there are grammar questions
        if grammar_questions:
            grammar_heading = doc.add_heading('Grammar Questions', level=1)
            grammar_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for heading_part in grammar_heading.runs:
                heading_part.font.size = Pt(format_options.get('heading_font_size', 14))
                heading_part.font.name = format_options.get('font_name', 'Calibri')
            
            start_index = len(vocab_questions) + 1 if vocab_questions else 1
            add_questions_to_document(doc, grammar_questions, format_options, start_index=start_index)
    else:
        # Add all questions without grouping
        add_questions_to_document(doc, exam_questions, format_options)
    
    # Add answer key if requested
    if include_answers:
        doc.add_page_break()
        answer_heading = doc.add_heading('Answer Key', level=1)
        answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for heading_part in answer_heading.runs:
            heading_part.font.size = Pt(format_options.get('heading_font_size', 14))
            heading_part.font.name = format_options.get('font_name', 'Calibri')
        
        # Create a table for answers (5 columns)
        num_questions = len(exam_questions)
        num_rows = (num_questions + 4) // 5  # Ceiling division
        
        if num_questions > 0:
            answer_table = doc.add_table(rows=num_rows, cols=5)
            answer_table.style = 'Table Grid'
            
            question_index = 0
            for row in range(num_rows):
                for col in range(5):
                    if question_index < num_questions:
                        cell = answer_table.cell(row, col)
                        correct_idx = exam_questions[question_index]['correct_answer']
                        correct_letter = chr(65 + correct_idx)  # A, B, C, D
                        
                        cell_paragraph = cell.paragraphs[0]
                        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_run = cell_paragraph.add_run(f"{question_index+1}. {correct_letter}")
                        cell_run.font.size = Pt(format_options.get('font_size', 11))
                        cell_run.font.name = format_options.get('font_name', 'Calibri')
                        
                        question_index += 1
    
    # Save document
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'exam.docx')
    doc.save(output_path)
    
    return output_path

def add_questions_to_document(doc, questions, format_options, start_index=1):
    """Add questions to the document with proper formatting"""
    for i, question in enumerate(questions):
        # Question number and text
        question_index = i + start_index
        q_paragraph = doc.add_paragraph()
        q_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        q_paragraph.paragraph_format.space_after = Pt(2)
        q_run = q_paragraph.add_run(f"{question_index}. {question['body']}")
        q_run.font.size = Pt(format_options.get('font_size', 11))
        q_run.font.name = format_options.get('font_name', 'Calibri')
        q_run.bold = True
        
        # Answer choices on a single line
        options_paragraph = doc.add_paragraph()
        options_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        options_paragraph.paragraph_format.left_indent = Inches(0.25)
        
        options_parts = []
        for j, answer in enumerate(question['answers']):
            option_letter = chr(65 + j)
            options_parts.append(f"{option_letter}. {answer}")
        
        options_text = "   ".join(options_parts)
        options_run = options_paragraph.add_run(options_text)
        options_run.font.size = Pt(format_options.get('font_size', 11))
        options_run.font.name = format_options.get('font_name', 'Calibri')
        
        # Add space between questions
        spacing = options_paragraph
        spacing.paragraph_format.space_after = Pt(format_options.get('question_spacing', 6))

def add_header_footer(doc, title):
    """Add header and footer to the document"""
    # Add header
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_paragraph.add_run(title)
    header_run.font.size = Pt(9)
    header_run.font.name = 'Calibri'
    
    # Add footer with page numbers
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    add_page_number(footer_paragraph)

def add_page_number(paragraph):
    """Add page number field to the paragraph"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
